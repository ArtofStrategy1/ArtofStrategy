# /home/elijah/data2int/backend/analysis-libraries/analysis-libraries-app/save_modules/save_docx.py

import io
import pandas as pd
import pygraphviz as pgv
import base64
from docx import Document
from docx.shared import Inches
from typing import Dict, Any, List

# --- HELPER FUNCTIONS ---

def _add_heading_docx(document: Document, item: Dict[str, Any]):
    level = item.get("level", 1)
    text = str(item.get("data", "Unnamed Heading"))
    document.add_heading(text, level=level)

def _add_paragraph_docx(document: Document, item: Dict[str, Any]):
    text = str(item.get("data", ""))
    document.add_paragraph(text)

def _add_list_item_docx(document: Document, item: Dict[str, Any]):
    text = str(item.get("data", ""))
    document.add_paragraph(text, style='List Bullet')

def _add_csv_table_docx(document: Document, item: Dict[str, Any]):
    csv_string = item.get("data", "")
    if not csv_string:
        return
        
    try:
        df = pd.read_csv(io.StringIO(csv_string))
        df = df.fillna('NA')

        table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        
        # Header row
        for j, col_name in enumerate(df.columns):
            table.cell(0, j).text = str(col_name)
            
        # Data rows
        for i, row in df.iterrows():
            for j, val in enumerate(row):
                table.cell(i + 1, j).text = str(val)
        
        document.add_paragraph() # Add space after table
            
    except Exception as e:
        document.add_paragraph(f"Error rendering table: {e}")

def _add_dot_image_docx(document: Document, item: Dict[str, Any]):
    dot_string = item.get("data", "")
    if not dot_string:
        return
        
    try:
        G = pgv.AGraph(string=dot_string)
        G.layout(prog="dot")
        img_bytes = G.draw(format="png")
        img_io = io.BytesIO(img_bytes)
        
        document.add_picture(img_io, width=Inches(6.0))
        document.add_paragraph()
        
    except Exception as e:
        document.add_paragraph(f"Error rendering DOT diagram. Is graphviz installed?\nError: {e}")

def _add_base64_image_docx(document: Document, item: Dict[str, Any]):
    b64_string = item.get("data", "")
    if not b64_string:
        return
        
    try:
        img_bytes = base64.b64decode(b64_string)
        img_io = io.BytesIO(img_bytes)
        
        document.add_picture(img_io, width=Inches(6.0))
        document.add_paragraph()
        
    except Exception as e:
        document.add_paragraph(f"Error rendering Base64 image: {e}")

# --- MAIN FUNCTION ---

# Map item types to the functions that render them
RENDER_MAP_DOCX = {
    "heading": _add_heading_docx,
    "paragraph": _add_paragraph_docx,
    "list_item": _add_list_item_docx,
    "table_csv": _add_csv_table_docx,
    "image_dot": _add_dot_image_docx,
    "image_base64": _add_base64_image_docx,
}

def create_docx_report(data: Dict[Any, Any]) -> bytes:
    """
    Generates a DOCX report by looping through a 'report_content' list.
    """
    document = Document()
    document.add_heading("S.A.G.E. Analysis Report", 0)

    report_content: List[Dict[str, Any]] = data.get("report_content", [])
    
    if not report_content:
        document.add_paragraph("No exportable content was provided.")
        
    for item in report_content:
        item_type = item.get("type")
        render_function = RENDER_MAP_DOCX.get(item_type)
        
        if render_function:
            try:
                render_function(document, item)
            except Exception as e:
                document.add_paragraph(f"--- CRITICAL ERROR RENDERING ITEM '{item_type}' ---\n{e}")
        else:
            p = document.add_paragraph()
            p.add_run(f"Unknown content type: '{item_type}'").italic = True

    # Save to a byte stream and return
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io.read()